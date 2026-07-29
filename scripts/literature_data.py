"""Parse literature values (priors and defaults) for 3-PG model parameters.

Builds per-species parameter tables from two sources: the docx supplement of
Trotsiuk et al. (posterior estimates and prior ranges for Picea abies and
Fagus sylvatica) and the Forrester et al. (2021) PDF (prior ranges and
parameter sources/defaults for twelve European species,
https://link.springer.com/article/10.1007/s10342-021-01370-3).
"""

import os
from typing import Literal

import camelot
import polars as pl
from docx import Document
from docx.table import Table

from trunx.config import project_root, threepg_data_folder
from trunx.gp3.model_inputs import Params

LITERATURE_FOLDER = os.path.join(project_root, "literature")
FORRESTER_PDF_PATH = os.path.join(LITERATURE_FOLDER, "s10342-021-01370-3.pdf")
TROTSIUK_DOCX_PATH = os.path.join(LITERATURE_FOLDER, "gcb15011-sup-0001-supinfo.docx")
SOLLING_RDA_PATH = os.path.join(
    project_root, "models/r3PG/vignettes_build/vignette_data/solling.rda"
)

# Some (parameter, species) pairs have no Table 2 prior range at all (e.g.
# gammaF1 for deciduous species) but do have a Table 4 "#" posterior
# estimate. "forrester" keeps that species-specific Forrester value;
# "default" instead falls back to param_default's single generic value.
FILL_NON_PRIOR_SOURCE: Literal["forrester", "default"] = "default"

FORRESTER_SPECIES = [
    "Abies alba",
    "Acer pseudoplatanus",
    "Betula pendula",
    "Fagus sylvatica",
    "Fraxinus excelsior",
    "Larix decidua",
    "Picea abies",
    "Pinus cembra",
    "Pinus sylvestris",
    "Pseudotsuga menziesii",
    "Quercus petraea",
    "Quercus robur",
]

# Table 2 — the 18 Bayesian-calibrated 3-PG parameters (prior min-max ranges).
TABLE2_PARAMETERS = [
    "pFS2",
    "pFS20",
    "pRx",
    "pRn",
    "gammaF1",
    "gammaR",
    "Tmin",
    "Topt",
    "Tmax",
    "fCalpha700",
    "fCg700",
    "wSx1000",
    "thinPower",
    "k",
    "MaxIntcptn",
    "alphaCx",
    "MaxCond",
    "CoeffCond",
]

# Table 4 parameters marked "†": defaults from Forrester and Tang (2016) /
# Sands and Landsberg (2002), shared across all species (plus fracBB0/
# fracBB1/tBB, which carry no source marker but are likewise fixed at 0).
DEFAULT_PARAMETERS = [
    "leafgrow",
    "leaffall",
    "kF",
    "m0",
    "fN0",
    "fNn",
    "nAge",
    "rAge",
    "gammaN1",
    "gammaN0",
    "tgammaN",
    "ngammaN",
    "fullCanAge",
    "LAImaxIntcptn",
    "cVPD",
    "Y",
    "MinCond",
    "LAIgcx",
    "BLcond",
    "RGcGw",
    "crownshape",
    "D13CTissueDif",
    "aFracDiffu",
    "bFracRubi",
    "fracBB0",
    "fracBB1",
    "tBB",
    "Qa",
    "Qb",
    "gDM_mol",
    "molPAR_MJ",
]

# Table 4 parameters marked "*": point values calculated from published
# studies (see Table S12 for sources), one value per species.
STAR_PARAMETERS = [
    "aWS",
    "nWS",
    "gammaF0",
    "tgammaF",
    "MaxAge",
    "SLA0",
    "SLA1",
    "tSLA",
    "rhoMin",
    "rhoMax",
    "tRho",
]

# Table 4 parameters marked "‡": point values calculated from EFM/FRN data
# (regression analyses), one value per species.
FRN_PARAMETERS = [
    "mF",
    "mR",
    "mS",
    "aH",
    "nHB",
    "nHC",
    "aV",
    "nVB",
    "nVH",
    "nVBH",
    "aK",
    "nKB",
    "nKH",
    "nKC",
    "nKrh",
    "aHL",
    "nHLB",
    "nHLL",
    "nHLC",
    "nHLrh",
]


def parse_prior_posterior_table(table: Table) -> pl.DataFrame:
    """Parse a Table S3/S4-style docx table (two-row header) into a DataFrame."""
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    columns = [
        "parameter",
        "prior_min",
        "prior_max",
        "posterior_2.5%",
        "posterior_50%",
        "posterior_97.5%",
        "reduction_pct",
        "conv_point_est",
        "conv_upper_ci",
    ]
    # The first two rows are the (merged) header; data starts on row 2.
    return pl.DataFrame(rows[2:], schema=columns, orient="row")


def load_trotsiuk_tables() -> tuple[pl.DataFrame, pl.DataFrame]:
    """Load the Picea abies and Fagus sylvatica prior/posterior docx tables.

    Returns
    -------
    tuple[pl.DataFrame, pl.DataFrame]
        Picea abies table, Fagus sylvatica table.
    """
    doc = Document(TROTSIUK_DOCX_PATH)
    return parse_prior_posterior_table(doc.tables[2]), parse_prior_posterior_table(doc.tables[3])


def build_species_param_df(
    species_table: pl.DataFrame,
    param_default: pl.DataFrame,
    species_name: str,
    params: list[str],
) -> pl.DataFrame:
    """Combine a species' docx prior/posterior table with the 3-PG defaults.

    Parameters
    ----------
    species_table : pl.DataFrame
        Parsed docx table for one species (see `parse_prior_posterior_table`).
    param_default : pl.DataFrame
        Default parameter values (`data.default.xlsx`).
    species_name : str
        Column name to give the species' posterior-median values.
    params : list[str]
        Parameter names to keep (`Params._fields`).

    Returns
    -------
    pl.DataFrame
        Columns: parameter, min, <species_name>, max.
    """
    species_df = species_table.select(
        "parameter", "prior_min", "posterior_50%", "prior_max"
    ).rename({"prior_min": "min", "posterior_50%": species_name, "prior_max": "max"})
    species_df = species_df.with_columns(
        pl.col("min", species_name, "max").cast(pl.Float64, strict=False)
    )
    species_df = param_default.join(species_df, on="parameter", how="left")
    species_df = species_df.with_columns(pl.col(species_name).fill_null(pl.col("default")))
    species_df = species_df.select("parameter", "min", species_name, "max")
    return species_df.filter(pl.col("parameter").is_in(params))


def _parse_range_token(token: str) -> tuple[float, float] | None:
    """Convert a "min–max" (or "*"/empty) token to a (min, max) tuple, or None."""
    token = token.strip()
    if token in ("", "*"):
        return None
    negative_first = token.startswith("−") or token.startswith("-")
    if negative_first:
        token = token[1:].strip()
    lo_str, hi_str = token.split("–")
    lo = -float(lo_str) if negative_first else float(lo_str)
    return lo, float(hi_str)


def _parse_point_estimate(token: str) -> float:
    """Extract the point estimate from a "value" or "value (lo–" token."""
    value = token.split("(")[0].strip()
    return float(value.replace("−", "-").replace(" ", ""))


def parse_forrester_prior_ranges(species: list[str]) -> pl.DataFrame:
    """Parse Table 2 (Bayesian-calibrated parameter prior ranges) from the Forrester PDF.

    Returns
    -------
    pl.DataFrame
        Columns: parameter, species, min, max.
    """
    rows = camelot.read_pdf(FORRESTER_PDF_PATH, pages="7", flavor="stream")[0].df.values.tolist()

    prior_rows = []
    i = 0
    while i < len(rows):
        name = rows[i][0].strip()
        if name in TABLE2_PARAMETERS:
            row1 = [c.strip() for c in rows[i][1 : 1 + len(species)]]
            is_continuation = i + 1 < len(rows) and rows[i + 1][0].strip() == ""
            row2 = (
                [c.strip() for c in rows[i + 1][1 : 1 + len(species)]]
                if is_continuation
                else [""] * len(species)
            )
            for sp, first, second in zip(species, row1, row2, strict=True):
                token = first + second if first.endswith("–") else first
                param_range = _parse_range_token(token)
                if param_range is None:
                    continue
                prior_rows.append(
                    {
                        "parameter": name,
                        "species": sp,
                        "min": param_range[0],
                        "max": param_range[1],
                    }
                )
            i += 2 if is_continuation else 1
        else:
            i += 1

    return pl.DataFrame(prior_rows)


def parse_forrester_posterior_defaults(
    table4_rows: list[list[list[str]]], species: list[str]
) -> pl.DataFrame:
    """Parse Table 4 parameters marked "#" (posterior median point estimates).

    Covers all species for each of the 18 Bayesian-calibrated parameters,
    including species with no Table 2 prior range (e.g. gammaF1 for
    deciduous species) — the point estimate always sits on the first of the
    three physical rows the CI wraps across.

    Returns
    -------
    pl.DataFrame
        Columns: parameter, species, default.
    """
    default_rows = []
    found = set()
    for rows in table4_rows:
        for row in rows:
            name = row[0].strip()
            if not name.endswith("#") or name[:-1] not in TABLE2_PARAMETERS:
                continue
            param = name[:-1]
            values = [c.strip() for c in row[1 : 1 + len(species)]]
            for sp, value in zip(species, values, strict=True):
                default_rows.append(
                    {"parameter": param, "species": sp, "default": _parse_point_estimate(value)}
                )
            found.add(param)

    missing = set(TABLE2_PARAMETERS) - found
    if missing:
        raise ValueError(f"Could not find posterior default rows for: {missing}")

    return pl.DataFrame(default_rows)


def load_table4_rows() -> list[list[list[str]]]:
    """Read Table 4 (parameter sources and defaults) from the Forrester PDF as raw rows."""
    tables = camelot.read_pdf(FORRESTER_PDF_PATH, pages="11-14", flavor="stream")
    return [table.df.values.tolist() for table in tables]


def parse_forrester_defaults(
    table4_rows: list[list[list[str]]], species: list[str]
) -> pl.DataFrame:
    """Parse Table 4 parameters marked "†" (defaults shared across all species).

    A few parameter names (LAImaxIntcptn, D13CTissueDif, molPAR_MJ) wrap
    across two rows; camelot keeps their values on the row carrying the
    first name fragment, so those are matched by prefix instead of an exact
    name match.

    Returns
    -------
    pl.DataFrame
        Columns: parameter, species, default.
    """

    def marker_for(param: str) -> str:
        return param if param in {"fracBB0", "fracBB1", "tBB"} else param + "†"

    default_rows = []
    found = set()
    for rows in table4_rows:
        for row in rows:
            name = row[0].strip()
            stripped = name.rstrip("-_")
            for param in DEFAULT_PARAMETERS:
                if param in found:
                    continue
                marker = marker_for(param)
                is_wrapped_name = stripped and stripped != name and marker.startswith(stripped)
                if name != marker and not is_wrapped_name:
                    continue
                values = [c.strip() for c in row[1:] if c.strip() != ""]
                if len(values) != len(species):
                    continue
                for sp, value in zip(species, values, strict=True):
                    default_rows.append(
                        {
                            "parameter": param,
                            "species": sp,
                            "default": _parse_point_estimate(value),
                        }
                    )
                found.add(param)
                break

    missing = set(DEFAULT_PARAMETERS) - found
    if missing:
        raise ValueError(f"Could not find default rows for: {missing}")

    return pl.DataFrame(default_rows)


def _parse_marked_point_values(
    table4_rows: list[list[list[str]]],
    species: list[str],
    marker: str,
    parameters: list[str],
) -> pl.DataFrame:
    """Parse Table 4 rows ending in `marker` into one point value per species.

    Shared by the "*" (published-study) and "‡" (EFM/FRN regression) marked
    parameters, which both have exactly one value per species and no range.

    Returns
    -------
    pl.DataFrame
        Columns: parameter, species, value.
    """
    value_rows = []
    found = set()
    for rows in table4_rows:
        for row in rows:
            name = row[0].strip()
            if not name.endswith(marker) or name[: -len(marker)] not in parameters:
                continue
            param = name[: -len(marker)]
            values = [c.strip() for c in row[1:] if c.strip() != ""]
            if len(values) != len(species):
                continue
            for sp, value in zip(species, values, strict=True):
                value_rows.append(
                    {"parameter": param, "species": sp, "value": _parse_point_estimate(value)}
                )
            found.add(param)

    missing = set(parameters) - found
    if missing:
        raise ValueError(f"Could not find rows marked {marker!r} for: {missing}")

    return pl.DataFrame(value_rows)


def parse_forrester_source_calculated(
    table4_rows: list[list[list[str]]], species: list[str]
) -> pl.DataFrame:
    """Parse Table 4 parameters marked "*" (point values calculated from published studies).

    Returns
    -------
    pl.DataFrame
        Columns: parameter, species, value.
    """
    return _parse_marked_point_values(table4_rows, species, "*", STAR_PARAMETERS)


def parse_forrester_frn_regression(
    table4_rows: list[list[list[str]]], species: list[str]
) -> pl.DataFrame:
    """Parse Table 4 parameters marked "‡" (point values from EFM/FRN regression).

    Returns
    -------
    pl.DataFrame
        Columns: parameter, species, value.
    """
    return _parse_marked_point_values(table4_rows, species, "‡", FRN_PARAMETERS)


def combine_parameter_tables(*tables: pl.DataFrame) -> pl.DataFrame:
    """Stack parameter tables into one, adding null min/max where absent.

    Parameters
    ----------
    *tables : pl.DataFrame
        Tables with columns parameter, species, default, and optionally
        min/max.

    Returns
    -------
    pl.DataFrame
        Columns: parameter, species, min, max, default.
    """
    columns = ["parameter", "species", "min", "max", "default"]
    aligned = [
        table.with_columns(
            *[
                pl.lit(None, dtype=pl.Float64).alias(col)
                for col in ("min", "max")
                if col not in table.columns
            ]
        ).select(columns)
        for table in tables
        if not table.is_empty()
    ]
    return pl.concat(aligned)


def fill_missing_gammaF1(
    forrester_prior_df: pl.DataFrame, param_default: pl.DataFrame, species: list[str]
) -> pl.DataFrame:
    """Fill in gammaF1 for species missing a Table 2 prior range.

    gammaF1 has no Table 2 prior range for deciduous species (marked "* Not
    applicable"), so joining Table 2 onto Table 4 drops those species
    entirely — fall back to param_default's generic value for them.

    Returns
    -------
    pl.DataFrame
        `forrester_prior_df` with any missing gammaF1 species added.
    """
    present = set(forrester_prior_df.filter(pl.col("parameter") == "gammaF1")["species"].to_list())
    missing_species = set(species) - present
    if not missing_species:
        return forrester_prior_df

    default_value = param_default.filter(pl.col("parameter") == "gammaF1")["default"].item()
    fill_df = pl.DataFrame(
        [
            {"parameter": "gammaF1", "species": sp, "default": default_value}
            for sp in missing_species
        ]
    )
    return combine_parameter_tables(forrester_prior_df, fill_df)


def fill_remaining_defaults(
    forrester_prior_df: pl.DataFrame,
    param_default: pl.DataFrame,
    params: list[str],
    species: list[str],
) -> pl.DataFrame:
    """Fill parameters not covered by `forrester_prior_df` from param_default.

    Broadcasts each missing parameter's generic default to every species,
    and prints which (if any) parameters needed this fallback.

    Returns
    -------
    pl.DataFrame
        `forrester_prior_df` combined with the param_default fallback rows.
    """
    missing = set(params) - set(forrester_prior_df["parameter"].unique().to_list())
    if missing:
        print(
            f"Parameters without a forrester literature source: {missing}, filling \
            from param_default"
        )
    else:
        print("All Params fields have a forrester literature source.")

    remaining_defaults = dict(
        param_default.filter(pl.col("parameter").is_in(list(missing)))
        .select("parameter", "default")
        .iter_rows()
    )
    remaining_default_df = pl.DataFrame(
        [
            {"parameter": param, "species": sp, "default": value}
            for param, value in remaining_defaults.items()
            for sp in species
        ]
    )
    return combine_parameter_tables(forrester_prior_df, remaining_default_df)


if __name__ == "__main__":
    params = list(Params._fields)
    param_default = pl.read_excel(os.path.join(threepg_data_folder, "data.default.xlsx"))
    piab_table, fasy_table = load_trotsiuk_tables()
    piab_df = build_species_param_df(piab_table, param_default, "Picea abies", params)
    fasy_df = build_species_param_df(fasy_table, param_default, "Fagus sylvatica", params)
    print(piab_df)
    print(fasy_df)

    table4_rows = load_table4_rows()
    forrester_min_max_df = parse_forrester_prior_ranges(FORRESTER_SPECIES)
    forrester_posterior_default_df = parse_forrester_posterior_defaults(
        table4_rows, FORRESTER_SPECIES
    )

    if FILL_NON_PRIOR_SOURCE == "forrester":
        forrester_prior_df = forrester_posterior_default_df.join(
            forrester_min_max_df, on=["parameter", "species"], how="left"
        )
        forrester_prior_df = combine_parameter_tables(
            forrester_prior_df,
            parse_forrester_defaults(table4_rows, FORRESTER_SPECIES),
            parse_forrester_source_calculated(table4_rows, FORRESTER_SPECIES).rename(
                {"value": "default"}
            ),
            parse_forrester_frn_regression(table4_rows, FORRESTER_SPECIES).rename(
                {"value": "default"}
            ),
        )
    else:
        forrester_prior_df = forrester_min_max_df.join(
            forrester_posterior_default_df, on=["parameter", "species"], how="left"
        )
        forrester_prior_df = fill_missing_gammaF1(
            forrester_prior_df, param_default, FORRESTER_SPECIES
        )

    param_df = fill_remaining_defaults(
        forrester_prior_df, param_default, params, FORRESTER_SPECIES
    )

    output_path = os.path.join(
        threepg_data_folder, f"literature_params_forrester_{FILL_NON_PRIOR_SOURCE}.parquet"
    )
    param_df.write_parquet(output_path)

    print(f"Saved param_df to {output_path}")

    print(param_df.height)
