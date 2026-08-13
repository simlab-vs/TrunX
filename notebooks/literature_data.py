"""Literature data processing for 3PG model implementation."""

import marimo

__generated_with = "0.23.8"
app = marimo.App(width="medium")


@app.cell
def _():
    import os

    import polars as pl
    import pyreadr

    from trunx.config import threepg_data_folder
    from trunx.gp3.model_inputs import Params

    return Params, os, pl, pyreadr, threepg_data_folder


@app.cell
def _(Params):
    params = list(Params._fields)
    return (params,)


@app.cell
def _(pyreadr):
    data = pyreadr.read_r("./models/r3PG/vignettes_build/vignette_data/solling.rda")
    data.keys()
    return


@app.cell
def _(os, pl, threepg_data_folder):
    from docx import Document

    param_default = pl.read_excel(os.path.join(threepg_data_folder, "data.default.xlsx"))

    def parse_prior_posterior_table(table):
        """Parse a Table S3/S4-style docx table (two-row header) into a DataFrame."""
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        columns = [
            "parameter",
            "prior_min",
            "prior_max",
            "posterior_2.5%",
            "posterior_50%",
            "posterior_97.5%",
            "reduction_pct",
            "conv_point_est",
            "conv_upper_ci",
        ]
        # The first two rows are the (merged) header; data starts on row 2.
        return pl.DataFrame(rows[2:], schema=columns, orient="row")

    file_path = os.path.join("./literature/", "gcb15011-sup-0001-supinfo.docx")
    doc = Document(file_path)
    piab_table = parse_prior_posterior_table(doc.tables[2])
    fasy_table = parse_prior_posterior_table(doc.tables[3])
    return fasy_table, param_default, piab_table


@app.cell
def _(param_default, params, piab_table, pl):
    # Species: P. abies
    piab_df = piab_table.select("parameter", "prior_min", "posterior_50%", "prior_max").rename(
        {"prior_min": "min", "posterior_50%": "Picea abies", "prior_max": "max"}
    )
    piab_df = piab_df.with_columns(
        pl.col("min", "Picea abies", "max").cast(pl.Float64, strict=False)
    )
    piab_df = param_default.join(piab_df, on="parameter", how="left")
    piab_df = piab_df.with_columns(pl.col("Picea abies").fill_null(pl.col("default")))
    piab_df = piab_df.select("parameter", "min", "Picea abies", "max")

    piab_df = piab_df.filter(pl.col("parameter").is_in(params))

    piab_df.head()
    return


@app.cell
def _(fasy_table, param_default, params, pl):
    # Species: F. sylvatica
    fasy_df = fasy_table.select("parameter", "prior_min", "posterior_50%", "prior_max").rename(
        {"prior_min": "min", "posterior_50%": "Fagus sylvatica", "prior_max": "max"}
    )
    fasy_df = fasy_df.with_columns(
        pl.col("min", "Fagus sylvatica", "max").cast(pl.Float64, strict=False)
    )
    fasy_df = param_default.join(fasy_df, on="parameter", how="left")
    fasy_df = fasy_df.with_columns(
        pl.col("Fagus sylvatica").fill_null(pl.col("default")).alias("fasy")
    )
    fasy_df = fasy_df.select("parameter", "min", "Fagus sylvatica", "max")

    fasy_df = fasy_df.filter(pl.col("parameter").is_in(params))
    return


@app.cell
def _():
    # Forrester et al. (2021) https://link.springer.com/article/10.1007/s10342-021-01370-3
    import camelot

    FORRESTER_PDF_PATH = "./literature/s10342-021-01370-3.pdf"

    FORRESTER_SPECIES = [
        "Abies alba",
        "Acer pseudoplatanus",
        "Betula pendula",
        "Fagus sylvatica",
        "Fraxinus excelsior",
        "Larix decidua",
        "Picea abies",
        "Pinus cembra",
        "Pinus sylvestris",
        "Pseudotsuga menziesii",
        "Quercus petraea",
        "Quercus robur",
    ]
    return FORRESTER_PDF_PATH, FORRESTER_SPECIES, camelot


@app.cell
def _(FORRESTER_PDF_PATH, FORRESTER_SPECIES, camelot, pl):
    # Table 2 — prior ranges (min-max) of the 18 Bayesian-calibrated 3-PG
    # parameters, per species ("*" marks gammaF1 as not applicable to the
    # deciduous species). camelot's "stream" flavor (whitespace-based, since
    # this table has no ruling lines) returns each row already split into
    # per-species cells; only the "min–" / "max" continuation across two
    # physical rows needs re-joining.
    _TABLE2_PARAMETERS = [
        "pFS2",
        "pFS20",
        "pRx",
        "pRn",
        "gammaF1",
        "gammaR",
        "Tmin",
        "Topt",
        "Tmax",
        "fCalpha700",
        "fCg700",
        "wSx1000",
        "thinPower",
        "k",
        "MaxIntcptn",
        "alphaCx",
        "MaxCond",
        "CoeffCond",
    ]

    def _parse_range_token(token):
        """Convert a "min–max" (or "*"/empty) token to a (min, max) tuple, or None."""
        token = token.strip()
        if token in ("", "*"):
            return None
        negative_first = token.startswith("−") or token.startswith("-")
        if negative_first:
            token = token[1:].strip()
        lo_str, hi_str = token.split("–")
        lo = -float(lo_str) if negative_first else float(lo_str)
        return lo, float(hi_str)

    _rows = camelot.read_pdf(FORRESTER_PDF_PATH, pages="7", flavor="stream")[0].df.values.tolist()

    _prior_rows = []
    _i = 0
    while _i < len(_rows):
        _name = _rows[_i][0].strip()
        if _name in _TABLE2_PARAMETERS:
            _row1 = [c.strip() for c in _rows[_i][1 : 1 + len(FORRESTER_SPECIES)]]
            _is_continuation = _i + 1 < len(_rows) and _rows[_i + 1][0].strip() == ""
            _row2 = (
                [c.strip() for c in _rows[_i + 1][1 : 1 + len(FORRESTER_SPECIES)]]
                if _is_continuation
                else [""] * len(FORRESTER_SPECIES)
            )
            for _species, _first, _second in zip(FORRESTER_SPECIES, _row1, _row2, strict=True):
                _token = _first + _second if _first.endswith("–") else _first
                _range = _parse_range_token(_token)
                if _range is None:
                    continue
                _prior_rows.append(
                    {
                        "parameter": _name,
                        "species": _species,
                        "min": _range[0],
                        "max": _range[1],
                    }
                )
            _i += 2 if _is_continuation else 1
        else:
            _i += 1

    forrester_prior_df = pl.DataFrame(_prior_rows)

    forrester_prior_df.head()
    return (forrester_prior_df,)


@app.cell
def _(FORRESTER_PDF_PATH, FORRESTER_SPECIES, camelot, pl):
    # Table 4 — parameters marked "†": defaults from Forrester and Tang
    # (2016) / Sands and Landsberg (2002), not calibrated per species (plus
    # fracBB0/fracBB1/tBB, which carry no source marker but are likewise
    # fixed at 0 for every species in this parameter set). A few parameter
    # names (LAImaxIntcptn, D13CTissueDif, molPAR_MJ) wrap across two rows;
    # camelot keeps their values on the row carrying the first name fragment,
    # so those are matched by prefix instead of requiring an exact name match.
    _DEFAULT_PARAMETERS = [
        "leafgrow",
        "leaffall",
        "kF",
        "m0",
        "fN0",
        "fNn",
        "nAge",
        "rAge",
        "gammaN1",
        "gammaN0",
        "tgammaN",
        "ngammaN",
        "fullCanAge",
        "LAImaxIntcptn",
        "cVPD",
        "Y",
        "MinCond",
        "LAIgcx",
        "BLcond",
        "RGcGw",
        "crownshape",
        "D13CTissueDif",
        "aFracDiffu",
        "bFracRubi",
        "fracBB0",
        "fracBB1",
        "tBB",
        "Qa",
        "Qb",
        "gDM_mol",
        "molPAR_MJ",
    ]

    def _marker_for(param):
        return param if param in {"fracBB0", "fracBB1", "tBB"} else param + "†"

    _tables = camelot.read_pdf(FORRESTER_PDF_PATH, pages="11-14", flavor="stream")

    _default_rows = []
    _found = set()
    for _table in _tables:
        for _row in _table.df.values.tolist():
            _name = _row[0].strip()
            _stripped = _name.rstrip("-_")
            for _param in _DEFAULT_PARAMETERS:
                if _param in _found:
                    continue
                _marker = _marker_for(_param)
                _is_wrapped_name = (
                    _stripped and _stripped != _name and _marker.startswith(_stripped)
                )
                if _name != _marker and not _is_wrapped_name:
                    continue
                _values = [c.strip() for c in _row[1:] if c.strip() != ""]
                if len(_values) != len(FORRESTER_SPECIES):
                    continue
                for _species, _value in zip(FORRESTER_SPECIES, _values, strict=True):
                    _default_rows.append(
                        {
                            "parameter": _param,
                            "species": _species,
                            "default": float(_value.replace("−", "-").replace(" ", "")),
                        }
                    )
                _found.add(_param)
                break

    _missing = set(_DEFAULT_PARAMETERS) - _found
    if _missing:
        raise ValueError(f"Could not find default rows for: {_missing}")

    forrester_default_df = pl.DataFrame(_default_rows)

    forrester_default_df.head()
    return (forrester_default_df,)


@app.cell
def _(FORRESTER_PDF_PATH, FORRESTER_SPECIES, camelot, pl):
    # Table 4 — parameters marked "*": calculated from published studies (see
    # Table S12 for sources), one point value per species (unlike the "#"
    # posterior parameters, these have no min-max quantile range).
    _STAR_PARAMETERS = [
        "aWS",
        "nWS",
        "gammaF0",
        "tgammaF",
        "MaxAge",
        "SLA0",
        "SLA1",
        "tSLA",
        "rhoMin",
        "rhoMax",
        "tRho",
    ]

    _tables = camelot.read_pdf(FORRESTER_PDF_PATH, pages="11-14", flavor="stream")

    _source_rows = []
    _found = set()
    for _table in _tables:
        for _row in _table.df.values.tolist():
            _name = _row[0].strip()
            if not _name.endswith("*") or _name[:-1] not in _STAR_PARAMETERS:
                continue
            _param = _name[:-1]
            _values = [c.strip() for c in _row[1:] if c.strip() != ""]
            if len(_values) != len(FORRESTER_SPECIES):
                continue
            for _species, _value in zip(FORRESTER_SPECIES, _values, strict=True):
                _source_rows.append(
                    {
                        "parameter": _param,
                        "species": _species,
                        "value": float(_value.replace("−", "-").replace(" ", "")),
                    }
                )
            _found.add(_param)

    _missing = set(_STAR_PARAMETERS) - _found
    if _missing:
        raise ValueError(f"Could not find source-calculated rows for: {_missing}")

    forrester_source_calculated_df = pl.DataFrame(_source_rows)

    forrester_source_calculated_df.head()
    return (forrester_source_calculated_df,)


@app.cell
def _(
    Params,
    forrester_default_df,
    forrester_prior_df,
    forrester_source_calculated_df,
):

    list_params = (
        forrester_prior_df["parameter"].unique().to_list()
        + forrester_default_df["parameter"].unique().to_list()
        + forrester_source_calculated_df["parameter"].unique().to_list()
    )

    full_params = Params._fields

    set(full_params) - set(list_params)
    return


@app.cell
def _(FORRESTER_PDF_PATH, FORRESTER_SPECIES, camelot, pl):
    # Table 4 — parameters marked "‡": calculated from EFM and FRN data
    # (regression analyses), one point value per species. This covers the
    # `set(full_params) - set(list_params)` remainder, minus SWconst/SWpower
    # which aren't in this paper's tables at all (sourced elsewhere, e.g.
    # Landsberg and Waring 1997).
    _FRN_PARAMETERS = [
        "mF",
        "mR",
        "mS",
        "aH",
        "nHB",
        "nHC",
        "aV",
        "nVB",
        "nVH",
        "nVBH",
        "aK",
        "nKB",
        "nKH",
        "nKC",
        "nKrh",
        "aHL",
        "nHLB",
        "nHLL",
        "nHLC",
        "nHLrh",
    ]

    _tables = camelot.read_pdf(FORRESTER_PDF_PATH, pages="11-14", flavor="stream")

    _frn_rows = []
    _found = set()
    for _table in _tables:
        for _row in _table.df.values.tolist():
            _name = _row[0].strip()
            if not _name.endswith("‡") or _name[:-1] not in _FRN_PARAMETERS:
                continue
            _param = _name[:-1]
            _values = [c.strip() for c in _row[1:] if c.strip() != ""]
            if len(_values) != len(FORRESTER_SPECIES):
                continue
            for _species, _value in zip(FORRESTER_SPECIES, _values, strict=True):
                _frn_rows.append(
                    {
                        "parameter": _param,
                        "species": _species,
                        "value": float(_value.replace("−", "-").replace(" ", "")),
                    }
                )
            _found.add(_param)

    _missing = set(_FRN_PARAMETERS) - _found
    if _missing:
        raise ValueError(f"Could not find EFM/FRN-regression rows for: {_missing}")

    forrester_frn_regression_df = pl.DataFrame(_frn_rows)

    forrester_frn_regression_df.head()
    return


@app.cell
def _(FORRESTER_SPECIES, os, pl, threepg_data_folder):
    # SWconst/SWpower aren't in Forrester et al. (2021) at all — take their
    # defaults from data.default.xlsx instead (Landsberg and Waring 1997),
    # broadcast to every species like the "†" defaults above.
    _defaults = pl.read_excel(os.path.join(threepg_data_folder, "data.default.xlsx"))
    _sw_defaults = dict(
        _defaults.filter(pl.col("parameter").is_in(["SWconst", "SWpower"]))
        .select("parameter", "default")
        .iter_rows()
    )

    forrester_sw_default_df = pl.DataFrame(
        [
            {"parameter": _param, "species": _species, "default": _value}
            for _param, _value in _sw_defaults.items()
            for _species in FORRESTER_SPECIES
        ]
    )

    forrester_sw_default_df.head()
    return


if __name__ == "__main__":
    app.run()
